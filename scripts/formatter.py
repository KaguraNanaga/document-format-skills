#!/usr/bin/env python3
"""
文档格式统一 v4
严格按照公文标准：
- 页边距：上37mm，下35mm，左28mm，右26mm
- 主标题：居中，二号（22pt），方正小标宋简体
- 正文：3号仿宋GB2312，首行缩进2字符，行距28磅
- 一级标题："一、" 三号黑体
- 二级标题："（一）" 三号楷体GB2312
- 三级标题："1." 三号仿宋GB2312
- 四级标题："（1）" 三号仿宋GB2312
"""

import sys
import re
from docx import Document
from docx.shared import Pt, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 字号对照：二号=22pt，三号=16pt，小四=12pt
# 2字符缩进 = 2 × 16pt = 32pt（三号字）

PRESETS = {
    'official': {
        'name': '公文格式',
        'page': {'top': 3.7, 'bottom': 3.5, 'left': 2.8, 'right': 2.6},
        # 主标题：二号方正小标宋简体，居中
        'title': {
            'font_cn': '方正小标宋简体',
            'font_en': 'Times New Roman',
            'size': 22,  # 二号
            'bold': False,
            'align': 'center',
            'indent': 0,
        },
        # 一级标题：三号黑体，"一、"，首行缩进2字符
        'heading1': {
            'font_cn': '黑体',
            'font_en': 'Times New Roman',
            'size': 16,  # 三号
            'bold': False,
            'align': 'left',
            'indent': 32,  # 2字符缩进
        },
        # 二级标题：三号楷体GB2312，"（一）"，首行缩进2字符
        'heading2': {
            'font_cn': '楷体_GB2312',
            'font_en': 'Times New Roman',
            'size': 16,
            'bold': False,
            'align': 'left',
            'indent': 32,
        },
        # 三级标题：三号仿宋GB2312，"1."，首行缩进2字符
        'heading3': {
            'font_cn': '仿宋_GB2312',
            'font_en': 'Times New Roman',
            'size': 16,
            'bold': False,
            'align': 'left',
            'indent': 32,
        },
        # 四级标题：三号仿宋GB2312，"（1）"，首行缩进2字符
        'heading4': {
            'font_cn': '仿宋_GB2312',
            'font_en': 'Times New Roman',
            'size': 16,
            'bold': False,
            'align': 'left',
            'indent': 32,
        },
        # 正文：三号仿宋GB2312，首行缩进2字符（32pt），行距28磅
        'body': {
            'font_cn': '仿宋_GB2312',
            'font_en': 'Times New Roman',
            'size': 16,
            'bold': False,
            'align': 'justify',
            'indent': 32,  # 2字符 = 2×16pt
            'line_spacing': 28,
        },
    },
    'academic': {
        'name': '学术论文格式',
        'page': {'top': 2.5, 'bottom': 2.5, 'left': 2.5, 'right': 2.5},
        'title': {'font_cn': '黑体', 'font_en': 'Times New Roman', 'size': 18, 'bold': True, 'align': 'center', 'indent': 0},
        'heading1': {'font_cn': '黑体', 'font_en': 'Times New Roman', 'size': 15, 'bold': True, 'align': 'left', 'indent': 0},
        'heading2': {'font_cn': '黑体', 'font_en': 'Times New Roman', 'size': 14, 'bold': True, 'align': 'left', 'indent': 0},
        'heading3': {'font_cn': '黑体', 'font_en': 'Times New Roman', 'size': 12, 'bold': False, 'align': 'left', 'indent': 0},
        'heading4': {'font_cn': '宋体', 'font_en': 'Times New Roman', 'size': 12, 'bold': False, 'align': 'left', 'indent': 0},
        'body': {'font_cn': '宋体', 'font_en': 'Times New Roman', 'size': 12, 'bold': False, 'align': 'justify', 'indent': 24, 'line_spacing': None},
    },
    'legal': {
        'name': '法律文书格式',
        'page': {'top': 3.0, 'bottom': 2.5, 'left': 3.0, 'right': 2.5},
        'title': {'font_cn': '宋体', 'font_en': 'Times New Roman', 'size': 22, 'bold': True, 'align': 'center', 'indent': 0},
        'heading1': {'font_cn': '黑体', 'font_en': 'Times New Roman', 'size': 14, 'bold': False, 'align': 'left', 'indent': 0},
        'heading2': {'font_cn': '黑体', 'font_en': 'Times New Roman', 'size': 14, 'bold': False, 'align': 'left', 'indent': 0},
        'heading3': {'font_cn': '宋体', 'font_en': 'Times New Roman', 'size': 14, 'bold': False, 'align': 'left', 'indent': 0},
        'heading4': {'font_cn': '宋体', 'font_en': 'Times New Roman', 'size': 14, 'bold': False, 'align': 'left', 'indent': 0},
        'body': {'font_cn': '宋体', 'font_en': 'Times New Roman', 'size': 14, 'bold': False, 'align': 'justify', 'indent': 28, 'line_spacing': None},
    },
}


def remove_background(doc):
    """移除页面背景颜色"""
    body = doc._body._body
    document = body.getparent()
    for elem in list(document):
        tag_name = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
        if tag_name == 'background':
            document.remove(elem)
    
    for para in doc.paragraphs:
        pPr = para._p.get_or_add_pPr()
        shd = pPr.find(qn('w:shd'))
        if shd is not None:
            pPr.remove(shd)
        for run in para.runs:
            run.font.highlight_color = None
            rPr = run._r.get_or_add_rPr()
            shd = rPr.find(qn('w:shd'))
            if shd is not None:
                rPr.remove(shd)


def detect_para_type(text, index, total, alignment):
    """
    检测段落类型
    返回: 'title', 'heading1', 'heading2', 'heading3', 'heading4', 'body'
    """
    text = text.strip()
    if not text:
        return 'empty'
    
    # 一级标题："一、" "二、" 等
    if re.match(r'^[一二三四五六七八九十]+、', text):
        return 'heading1'
    
    # 二级标题："（一）" "（二）" 等
    if re.match(r'^（[一二三四五六七八九十]+）', text):
        return 'heading2'
    if re.match(r'^\([一二三四五六七八九十]+\)', text):
        return 'heading2'
    
    # 三级标题："1." "2." 等（注意后面要有内容，且整体较短才是标题）
    if re.match(r'^\d+\.\s*\S', text) and len(text) < 60:
        return 'heading3'
    
    # 四级标题："（1）" "（2）" 等
    if re.match(r'^（\d+）', text) and len(text) < 60:
        return 'heading4'
    if re.match(r'^\(\d+\)', text) and len(text) < 60:
        return 'heading4'
    
    # 主标题判断（只在前3段，且满足条件）
    if index < 3:
        # 明确的标题模式
        if re.match(r'^关于.{2,30}的(通知|报告|请示|函|意见|决定|公告|通报|批复|汇报|方案|总结)$', text):
            return 'title'
        if re.match(r'^.{2,20}(通知|报告|请示|函|意见|决定|公告|通报|批复|汇报材料|工作汇报|工作方案|工作总结)$', text):
            return 'title'
        # 居中的短文本
        if alignment == WD_ALIGN_PARAGRAPH.CENTER and len(text) < 50:
            return 'title'
    
    # 其他都是正文
    return 'body'


def set_font(run, font_cn, font_en, size, bold=False):
    """设置字体"""
    run.font.name = font_en
    run.font.size = Pt(size)
    run.font.bold = bold
    
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_cn)


def format_paragraph(para, fmt, para_type, line_spacing_pt=28):
    """格式化段落"""
    pf = para.paragraph_format
    
    # 对齐方式
    align_map = {
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'left': WD_ALIGN_PARAGRAPH.LEFT,
        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    pf.alignment = align_map.get(fmt.get('align', 'justify'), WD_ALIGN_PARAGRAPH.JUSTIFY)
    
    # 首行缩进
    indent = fmt.get('indent', 0)
    if indent > 0:
        pf.first_line_indent = Pt(indent)
    else:
        pf.first_line_indent = Pt(0)
    
    # 行距（固定值28磅）
    ls = fmt.get('line_spacing', line_spacing_pt)
    if ls:
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(ls)
    else:
        pf.line_spacing = 1.5
    
    # 段前段后
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    
    # 字体
    for run in para.runs:
        set_font(run, fmt['font_cn'], fmt['font_en'], fmt['size'], fmt.get('bold', False))


def add_page_number(doc):
    """添加页码（底端居中）"""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        
        # 清空现有内容
        for para in footer.paragraphs:
            para.clear()
        
        # 添加页码
        if footer.paragraphs:
            para = footer.paragraphs[0]
        else:
            para = footer.add_paragraph()
        
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加页码域
        run = para.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar1)
        
        run2 = para.add_run()
        instrText = OxmlElement('w:instrText')
        instrText.text = 'PAGE'
        run2._r.append(instrText)
        
        run3 = para.add_run()
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run3._r.append(fldChar2)


def format_document(input_path, output_path, preset_name='official'):
    """格式化文档"""
    if preset_name not in PRESETS:
        print(f'Unknown preset: {preset_name}')
        print(f'Available: {", ".join(PRESETS.keys())}')
        sys.exit(1)
    
    preset = PRESETS[preset_name]
    print(f'Preset: {preset["name"]}')
    print(f'Input: {input_path}')
    
    doc = Document(input_path)
    total_paras = len(doc.paragraphs)
    
    # 1. 移除背景
    print('1. Removing background...')
    remove_background(doc)
    
    # 2. 设置页面边距
    print('2. Setting page margins...')
    page = preset['page']
    for section in doc.sections:
        section.top_margin = Cm(page['top'])
        section.bottom_margin = Cm(page['bottom'])
        section.left_margin = Cm(page['left'])
        section.right_margin = Cm(page['right'])
    
    # 3. 格式化段落
    print('3. Formatting paragraphs...')
    stats = {'title': 0, 'heading1': 0, 'heading2': 0, 'heading3': 0, 'heading4': 0, 'body': 0}
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        
        para_type = detect_para_type(text, i, total_paras, para.paragraph_format.alignment)
        
        # 选择对应的格式
        fmt_key = para_type if para_type in preset else 'body'
        fmt = preset.get(fmt_key, preset['body'])
        
        format_paragraph(para, fmt, para_type)
        stats[para_type] = stats.get(para_type, 0) + 1
        
        # 打印处理信息
        preview = text[:35] + '...' if len(text) > 35 else text
        print(f'   [{para_type:8}] {preview}')
    
    # 4. 处理表格
    print('4. Formatting tables...')
    body_fmt = preset['body']
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        for run in para.runs:
                            set_font(run, body_fmt['font_cn'], body_fmt['font_en'], body_fmt['size'])
                        para.paragraph_format.first_line_indent = Pt(0)
                        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                        para.paragraph_format.line_spacing = Pt(28)
    
    # 5. 添加页码
    print('5. Adding page numbers...')
    add_page_number(doc)
    
    # 保存
    doc.save(output_path)
    
    print()
    print('=' * 50)
    print('Statistics:')
    for k, v in stats.items():
        if v > 0:
            print(f'  {k}: {v}')
    print(f'Output: {output_path}')


if __name__ == '__main__':
    if len(sys.argv) < 3:
        print('Usage: python formatter.py input.docx output.docx [--preset official|academic|legal]')
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2]
    
    preset = 'official'
    if '--preset' in sys.argv:
        idx = sys.argv.index('--preset')
        if idx + 1 < len(sys.argv):
            preset = sys.argv[idx + 1]
    
    format_document(input_file, output_file, preset)
