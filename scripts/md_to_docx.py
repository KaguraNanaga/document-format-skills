#!/usr/bin/env python3
"""Convert Markdown content to DOCX and optionally apply formatting presets.

Core design:
- H1 becomes centered 22pt title paragraph, followed by a blank line.
- ## / #(section) -> 一、 (Heading 1 style, Word ToC visible)
- ### -> （一） (Heading 2 style)
- #### -> 1. (Heading 3 style)
- Inline **bold** is split into separate runs within the same paragraph.
- Hierarchical counters auto-reset under each parent heading.
- Final diagnostics run by default for every conversion.
"""

import argparse
import re
import shutil
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from analyzer import analyze_font, analyze_numbering, analyze_paragraph_format, analyze_punctuation, print_report
from formatter import PRESETS, format_document
from punctuation import process_document

CN_NUMS = [
    '一', '二', '三', '四', '五', '六', '七', '八', '九', '十',
    '十一', '十二', '十三', '十四', '十五', '十六', '十七', '十八', '十九', '二十',
]


def _cn(n):
    return CN_NUMS[n - 1] if 1 <= n <= len(CN_NUMS) else str(n)


def unique_path(path):
    path = Path(path)
    if not path.exists():
        return path
    for idx in range(2, 1000):
        candidate = path.with_name(f"{path.stem}_v{idx}{path.suffix}")
        if not candidate.exists():
            return candidate
    raise RuntimeError(f"Cannot find available output path for {path}")


def _parse_markdown_inline(text):
    """Parse inline Markdown, returning [(text, is_bold), ...]."""
    parts = []
    pattern = re.compile(r'\*\*([^*\n]+)\*\*|__([^_\n]+)__')
    last_end = 0
    for m in pattern.finditer(text):
        if m.start() > last_end:
            parts.append((text[last_end:m.start()], False))
        bold_text = m.group(1) or m.group(2)
        if bold_text:
            parts.append((bold_text, True))
        last_end = m.end()
    if last_end < len(text):
        parts.append((text[last_end:], False))
    return parts if parts else [(text, False)]


def _add_para_with_inline(doc, text, alignment=None, style=None):
    para = doc.add_paragraph()
    if style:
        para.style = doc.styles[style]
    if alignment is not None:
        para.alignment = alignment
    for content, is_bold in _parse_markdown_inline(text):
        if not content:
            continue
        run = para.add_run(content)
        if is_bold:
            run.font.bold = True
    return para


def _add_blank_paragraph(doc):
    """Add a blank paragraph with standard body line spacing."""
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing = Pt(28)
    return para


def convert_markdown_to_docx(input_path, output_path=None, title=None, overwrite=False):
    input_path = Path(input_path)
    if output_path is None:
        output_path = input_path.with_suffix(".docx")
    output_path = Path(output_path)
    if output_path.exists() and not overwrite:
        output_path = unique_path(output_path)

    md_text = input_path.read_text(encoding="utf-8")
    lines = md_text.split("\n")

    doc = Document()

    h2_counter = 0
    h3_counters = {}
    h4_counters = {}
    current_h2 = 0
    current_h3 = 0

    title_added = False
    in_code_block = False
    code_buffer = []

    h_pat = re.compile(r'^\s*(#{1,6})\s*(.*)$')
    ul_pat = re.compile(r'^\s*[-*+]\s+(.*)$')
    ol_pat = re.compile(r'^\s*\d+\.\s+(.*)$')
    quote_pat = re.compile(r'^>\s+(.*)$')

    first_h1_text = None
    for line in lines:
        if not line.strip():
            continue
        m = h_pat.match(line)
        if m and len(m.group(1)) == 1:
            first_h1_text = m.group(2).strip()
            break

    h1_is_section = bool(first_h1_text and re.match(r'^[一二三四五六七八九十]+[、，]', first_h1_text))

    for line in lines:
        if line.strip().startswith("```") or line.strip().startswith("~~~"):
            if in_code_block:
                for code_line in code_buffer:
                    if code_line.strip():
                        doc.add_paragraph(code_line)
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_buffer.append(line)
            continue

        if not line.strip():
            continue

        h_match = h_pat.match(line)
        if h_match:
            level = len(h_match.group(1))
            content = h_match.group(2).strip()

            if level == 1:
                if h1_is_section:
                    h2_counter += 1
                    current_h2 = h2_counter
                    _add_para_with_inline(doc, content, style='Heading 1')
                    continue
                if not title_added:
                    title_para = _add_para_with_inline(doc, title or content, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                    for run in title_para.runs:
                        run.font.size = Pt(22)
                    title_added = True
                    _add_blank_paragraph(doc)
                else:
                    h2_counter += 1
                    _add_para_with_inline(doc, f"{_cn(h2_counter)}、{content}", style='Heading 1')
                    current_h2 = h2_counter
                continue
            elif level == 2:
                h2_counter += 1
                current_h2 = h2_counter
                _add_para_with_inline(doc, f"{_cn(h2_counter)}、{content}", style='Heading 1')
                continue
            elif level == 3:
                h3_counters[current_h2] = h3_counters.get(current_h2, 0) + 1
                current_h3 = h3_counters[current_h2]
                _add_para_with_inline(doc, f"（{_cn(current_h3)}）{content}", style='Heading 2')
                continue
            elif level >= 4:
                key = (current_h2, current_h3)
                h4_counters[key] = h4_counters.get(key, 0) + 1
                _add_para_with_inline(doc, f"{h4_counters[key]}. {content}", style='Heading 3')
                continue

        ul_match = ul_pat.match(line)
        ol_match = ol_pat.match(line)
        if ul_match:
            _add_para_with_inline(doc, ul_match.group(1).strip())
            continue
        if ol_match:
            _add_para_with_inline(doc, ol_match.group(1).strip())
            continue

        q_match = quote_pat.match(line)
        if q_match:
            _add_para_with_inline(doc, q_match.group(1).strip())
            continue

        _add_para_with_inline(doc, line.strip())

    if in_code_block:
        for code_line in code_buffer:
            if code_line.strip():
                doc.add_paragraph(code_line)

    if not title_added and (title or input_path.stem):
        fallback = (title or input_path.stem).strip()
        title_para = _add_para_with_inline(doc, fallback, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        for run in title_para.runs:
            run.font.size = Pt(22)
        body = doc.element.body
        body.remove(title_para._p)
        body.insert(0, title_para._p)
        # Blank line after title
        blank = _add_blank_paragraph(doc)
        body.remove(blank._p)
        body.insert(1, blank._p)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


def default_output_path(input_path, preset=None):
    input_path = Path(input_path)
    suffix = f"_{preset}" if preset else ""
    return input_path.with_name(f"{input_path.stem}{suffix}.docx")


def diagnose(path):
    doc = Document(path)
    results = {
        "punctuation": analyze_punctuation(doc),
        "numbering": analyze_numbering(doc),
        "paragraph": analyze_paragraph_format(doc),
        "font": analyze_font(doc),
    }
    print()
    print("Final diagnostics:")
    print_report(results)
    return results


def remove_page_number_footers(path):
    """Remove footer page-number lines added by formatter.py."""
    doc = Document(path)
    for section in doc.sections:
        for footer in (section.footer, section.even_page_footer, section.first_page_footer):
            footer.is_linked_to_previous = False
            for paragraph in footer.paragraphs:
                paragraph.clear()
    doc.save(path)


def convert_markdown_pipeline(
    input_path,
    output_path=None,
    title=None,
    overwrite=False,
    preset=None,
    draft_only=False,
    no_punctuation=False,
    skip_diagnose=False,
    keep_temp=False,
    no_page_number=False,
):
    input_path = Path(input_path)
    if preset and preset not in PRESETS and preset != "custom":
        available = sorted(list(PRESETS.keys()) + ["custom"])
        raise ValueError(f"Unknown preset: {preset}. Available: {', '.join(available)}")

    if draft_only and preset:
        raise ValueError("--draft-only cannot be used with --preset")

    if output_path is None:
        output_path = default_output_path(input_path, preset if not draft_only else None)
    output_path = Path(output_path)
    if output_path.exists() and not overwrite:
        output_path = unique_path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    temp_root = output_path.parent if keep_temp else Path(tempfile.mkdtemp(prefix="md_docx_"))
    draft_path = temp_root / f"{input_path.stem}_draft.docx"
    punct_path = temp_root / f"{input_path.stem}_punctuation.docx"
    formatted_path = temp_root / f"{input_path.stem}_formatted.docx"

    try:
        print("1. Markdown -> DOCX draft")
        convert_markdown_to_docx(input_path, draft_path, title=title, overwrite=True)
        current_path = draft_path

        if not draft_only and not no_punctuation:
            print("2. Fix punctuation")
            process_document(str(current_path), str(punct_path))
            current_path = punct_path
        else:
            print("2. Skip punctuation")

        if preset and not draft_only:
            print(f"3. Apply {preset} format")
            format_document(str(current_path), str(formatted_path), preset)
            current_path = formatted_path
            if no_page_number:
                print("3.1. Remove page-number footer")
                remove_page_number_footers(current_path)
        else:
            print("3. Skip formatting preset")

        shutil.copyfile(current_path, output_path)

        if skip_diagnose:
            print("4. Skip diagnostics")
        else:
            print("4. Run diagnostics")
            diagnose(output_path)
    finally:
        if not keep_temp and temp_root.exists():
            shutil.rmtree(temp_root)

    return output_path


def main():
    parser = argparse.ArgumentParser(description="Convert Markdown to DOCX, with optional formatting presets and diagnostics.")
    parser.add_argument("input", help="Input .md file")
    parser.add_argument("output", nargs="?", help="Output .docx file")
    parser.add_argument("--title", help="Override Word main title")
    parser.add_argument("--overwrite", action="store_true", help="Overwrite output if it exists")
    parser.add_argument("--preset", choices=sorted(list(PRESETS.keys()) + ["custom"]), help="Apply a formatter preset after Markdown conversion")
    parser.add_argument("--draft-only", action="store_true", help="Only convert Markdown to DOCX, then run diagnostics")
    parser.add_argument("--no-punctuation", action="store_true", help="Skip punctuation fixing")
    parser.add_argument("--skip-diagnose", action="store_true", help="Skip final diagnostics")
    parser.add_argument("--keep-temp", action="store_true", help="Keep intermediate DOCX files next to output")
    parser.add_argument("--no-page-number", action="store_true", help="Remove page-number footer after formatting")
    args = parser.parse_args()
    output = convert_markdown_pipeline(
        args.input,
        args.output,
        title=args.title,
        overwrite=args.overwrite,
        preset=args.preset,
        draft_only=args.draft_only,
        no_punctuation=args.no_punctuation,
        skip_diagnose=args.skip_diagnose,
        keep_temp=args.keep_temp,
        no_page_number=args.no_page_number,
    )
    print(f"Done: {output}")


if __name__ == "__main__":
    main()
